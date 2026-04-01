"""Chargement de documents: PDF, DOCX, TXT, MD, HTML, URL."""
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import urllib.request
import urllib.parse


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]

def _load_excel(path: str) -> str:
    """Charge un fichier Excel et convertit en texte structuré."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Feuille: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            # Ignore les lignes complètement vides
            values = [str(v).strip() for v in row if v is not None]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)

def _load_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _load_docx(path: str) -> str:
    from docx import Document as DocxDoc
    doc = DocxDoc(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _load_html(path: str) -> str:
    from bs4 import BeautifulSoup
    with open(path, encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _load_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".doc": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".txt": _load_text,
    ".md": _load_text,
    ".markdown": _load_text,
    ".rst": _load_text,
    ".csv": _load_text,
    ".json": _load_text,
      ".xlsx":     _load_excel,   # ← ajouté
    ".xls":      _load_excel,   # ← ajouté
}


def load_document(path: str) -> Document:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in LOADERS:
        raise ValueError(f"Format non supporté: {ext}")

    content = LOADERS[ext](str(p))
    return Document(
        content=content,
        metadata={
            "source": str(p),
            "filename": p.name,
            "extension": ext,
            "size_bytes": p.stat().st_size,
        },
    )


def scrape_url(url: str, timeout: int = 15) -> Document:
    """Scrape une URL HTTP/HTTPS et retourne un Document."""
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL invalide (http/https requis): {url}")

    req = urllib.request.Request(url, headers={"User-Agent": "RAGBot/1.0"})
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        raw_html = resp.read()

    from bs4 import BeautifulSoup
    soup = BeautifulSoup(raw_html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    title = soup.title.string.strip() if soup.title and soup.title.string else parsed.netloc
    content = soup.get_text(separator="\n", strip=True)

    # Nom de fichier safe pour les IDs ChromaDB
    safe_name = urllib.parse.quote(url, safe="").replace("%", "_")[:80] + ".url"

    return Document(
        content=content,
        metadata={
            "source": url,
            "filename": safe_name,
            "extension": ".url",
            "title": title,
        },
    )


def load_directory(directory: str) -> List[Document]:
    docs = []
    for file in sorted(Path(directory).rglob("*")):
        if file.is_file() and file.suffix.lower() in LOADERS:
            try:
                doc = load_document(str(file))
                if doc.content.strip():
                    docs.append(doc)
                    print(f"  Chargé: {file.name} ({len(doc.content):,} chars)")
            except Exception as e:
                print(f"  Ignoré {file.name}: {e}")
    return docs
